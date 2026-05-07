# app/UI/app.py - PRODUCTION VERSION v4.2 - Complete Fixed

import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent.parent))

import os
import json
import streamlit as st
import torch
import pandas as pd
import plotly.graph_objects as go
from transformers import AutoTokenizer, AutoModelForSequenceClassification
import re
import time
import io
import datetime
import numpy as np

# Matplotlib - pure python diagram generation (no cairo needed)
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch

# File parsing
from docx import Document
import pdfplumber

# Report generation
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer,
    Image as RLImage
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# UML pipeline
from app.ml.uml.parser          import parse_requirements
from app.ml.uml.graph_builder   import build_graph
from app.ml.uml.diagram_renderer import draw_uml_diagram_png

# Validity detection
from app.ml.validity.predict    import predict_validity
# ──────────────────────────────────────────────────────────────────────────────
# PAGE CONFIG
# ──────────────────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="AI-RADA | Requirement Analysis & Design Assistant",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ──────────────────────────────────────────────────────────────────────────────
# CSS
# ──────────────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
:root {
    --primary:      #1e88e5;
    --primary-dark: #1565c0;
    --success:      #43a047;
    --warning:      #f57c00;
    --danger:       #e53935;
    --light-bg:     #f5f7fa;
    --light-border: #e0e7ff;
    --text-dark:    #1a237e;
    --text-light:   #455a64;
    --shadow:       rgba(0,0,0,0.08);
}
* { margin:0; padding:0; box-sizing:border-box; }
html, body, [data-testid="stAppViewContainer"] {
    background:#f8fafb;
    font-family:'Segoe UI',Roboto,Tahoma,sans-serif;
    color:var(--text-dark);
}
[data-testid="stAppViewContainer"] > .main {
    background:#f8fafb; padding-top:0 !important;
}
[data-testid="stHeader"]  { background:transparent; border:none; }
[data-testid="stToolbar"] { display:none; }
#MainMenu, footer         { visibility:hidden; }

/* Hero */
.hero {
    background:linear-gradient(135deg,#1e88e5 0%,#1565c0 50%,#0d47a1 100%);
    padding:4rem 2rem 3rem; text-align:center;
    margin:-1rem -1rem 2.5rem -1rem;
    box-shadow:0 8px 24px rgba(30,136,229,0.2);
    position:relative; overflow:hidden;
}
.hero::before {
    content:''; position:absolute; inset:0;
    background:radial-gradient(ellipse at 30% 50%,rgba(255,255,255,0.12) 0%,transparent 70%);
    pointer-events:none;
}
.hero-title {
    font-size:3.5rem; font-weight:900; color:white;
    letter-spacing:2px; margin-bottom:0.5rem;
    text-shadow:0 2px 8px rgba(0,0,0,0.15); position:relative; z-index:1;
}
.hero-sub {
    font-size:1.2rem; color:rgba(255,255,255,0.95);
    margin-bottom:1rem; position:relative; z-index:1;
}
.hero-badges {
    display:flex; justify-content:center; gap:1rem;
    flex-wrap:wrap; margin-top:1rem; position:relative; z-index:1;
}
.badge {
    background:rgba(255,255,255,0.2);
    border:1px solid rgba(255,255,255,0.4);
    border-radius:25px; padding:0.4rem 1.2rem;
    font-size:0.9rem; font-weight:600; color:white;
    backdrop-filter:blur(8px);
}

/* Panel */
.panel {
    background:white; border:1px solid var(--light-border);
    border-radius:12px; padding:2rem; margin-bottom:1.8rem;
    box-shadow:0 2px 8px var(--shadow); transition:all 0.3s ease;
}
.panel:hover {
    border-color:var(--primary);
    box-shadow:0 4px 12px rgba(30,136,229,0.15);
}
.panel-title {
    font-size:1.25rem; font-weight:700; color:var(--primary);
    margin-bottom:1.2rem; display:flex; align-items:center; gap:0.7rem;
    padding-bottom:0.75rem; border-bottom:2px solid var(--light-border);
}

/* Section heading */
.section-heading {
    font-size:1.6rem; font-weight:800; color:var(--primary);
    margin:1rem 0 1.5rem; padding-bottom:1rem;
    border-bottom:3px solid var(--primary);
}

/* Metric grid */
.metric-grid {
    display:grid;
    grid-template-columns:repeat(auto-fit,minmax(170px,1fr));
    gap:1.2rem; margin-bottom:1.5rem;
}
.metric-card {
    background:linear-gradient(135deg,#f5f7fa 0%,#e8f0fe 100%);
    border:2px solid var(--light-border); border-radius:12px;
    padding:1.5rem; text-align:center; transition:all 0.3s ease;
}
.metric-card:hover {
    transform:translateY(-4px);
    box-shadow:0 6px 16px rgba(30,136,229,0.2);
    border-color:var(--primary);
}
.metric-value {
    font-size:2.8rem; font-weight:900; color:var(--primary);
    line-height:1; margin-bottom:0.5rem;
}
.metric-label {
    font-size:0.85rem; color:var(--text-light);
    text-transform:uppercase; letter-spacing:1px; font-weight:600;
}
.metric-card.green  { border-color:var(--success); background:linear-gradient(135deg,#f1f8f5,#e8f5f3); }
.metric-card.green .metric-value  { color:var(--success); }
.metric-card.orange { border-color:var(--warning); background:linear-gradient(135deg,#fff3f0,#ffe0d5); }
.metric-card.orange .metric-value { color:var(--warning); }
.metric-card.red    { border-color:var(--danger);  background:linear-gradient(135deg,#ffebee,#ffcdd2); }
.metric-card.red .metric-value    { color:var(--danger); }

/* Module cards */
.module-card {
    background:white; border:2px solid var(--light-border);
    border-radius:12px; padding:1.5rem; text-align:center;
    cursor:pointer; transition:all 0.3s ease; height:100%;
}
.module-card:hover {
    border-color:var(--primary);
    background:linear-gradient(135deg,#f5f7fa,#e8f0fe);
    transform:translateY(-3px);
    box-shadow:0 6px 16px rgba(30,136,229,0.2);
}
.module-icon  { font-size:2.5rem; margin-bottom:0.7rem; display:block; }
.module-name  { font-size:1rem; font-weight:700; color:var(--text-dark); margin-bottom:0.3rem; }
.module-desc  { font-size:0.8rem; color:var(--text-light); line-height:1.4; }

/* Buttons */
.stButton > button[kind="primary"] {
    background:linear-gradient(135deg,var(--primary),var(--primary-dark)) !important;
    color:white !important; border:none !important; border-radius:8px !important;
    font-weight:700 !important; font-size:1.05rem !important;
    padding:0.8rem 2rem !important;
    box-shadow:0 4px 12px rgba(30,136,229,0.3) !important;
    transition:all 0.3s ease !important;
}
.stButton > button[kind="primary"]:hover {
    transform:translateY(-2px) !important;
    box-shadow:0 6px 20px rgba(30,136,229,0.4) !important;
}

/* DataFrame */
[data-testid="stDataFrame"] {
    border-radius:10px !important;
    overflow:hidden !important;
    border:1px solid var(--light-border) !important;
}

/* Expander */
[data-testid="stExpander"] {
    background:white !important;
    border:1px solid var(--light-border) !important;
    border-radius:10px !important;
}

/* Tabs */
[data-testid="stTabs"] [data-baseweb="tab-list"] {
    background:#f5f7fa;
    border-bottom:2px solid var(--light-border); gap:0;
}
[data-testid="stTabs"] [data-baseweb="tab"] {
    color:var(--text-light); font-weight:600;
    padding:1rem 1.5rem !important; border:none !important;
}
[data-testid="stTabs"] [aria-selected="true"] {
    color:var(--primary) !important;
    border-bottom:3px solid var(--primary) !important;
}

/* Pills */
.pill { display:inline-block; padding:0.3rem 0.8rem; border-radius:20px; font-size:0.8rem; font-weight:700; margin:0.2rem; }
.pill-blue   { background:#e3f2fd; color:var(--primary); }
.pill-green  { background:#e8f5e9; color:var(--success); }
.pill-orange { background:#fff3e0; color:var(--warning); }
.pill-purple { background:#f3e5f5; color:#7b1fa2; }

/* UML panel */
.uml-panel {
    background:#f5f7fa; border:2px solid var(--light-border);
    border-radius:12px; padding:2rem; margin-top:1.5rem;
}
.uml-title {
    font-size:1.6rem; font-weight:800; color:var(--primary);
    margin-bottom:1.2rem; padding-bottom:0.75rem;
    border-bottom:2px solid var(--light-border);
}

/* Download box */
.dl-box {
    background:white; border:1px solid var(--light-border);
    border-radius:10px; padding:1.2rem; margin-bottom:0.8rem;
    text-align:center;
}
.dl-box-title { font-size:0.9rem; font-weight:700; color:var(--text-dark); margin-bottom:0.3rem; }
.dl-box-desc  { font-size:0.75rem; color:var(--text-light); margin-bottom:0.8rem; }

/* Progress */
[data-testid="stProgress"] > div > div {
    background:linear-gradient(90deg,var(--primary),#1976d2) !important;
    border-radius:4px !important;
}

/* Scrollbar */
::-webkit-scrollbar { width:8px; height:8px; }
::-webkit-scrollbar-track { background:#f5f7fa; }
::-webkit-scrollbar-thumb { background:#cbd5e0; border-radius:4px; }
::-webkit-scrollbar-thumb:hover { background:#a0aec0; }

/* Footer */
.footer {
    text-align:center; color:var(--text-light); font-size:0.9rem;
    margin-top:4rem; padding:2rem 1rem;
    border-top:1px solid var(--light-border);
}
.footer strong { color:var(--primary); }
</style>
""", unsafe_allow_html=True)

# ──────────────────────────────────────────────────────────────────────────────
# SESSION STATE
# ──────────────────────────────────────────────────────────────────────────────
_defaults = {
    "analysis_results":  None,
    "uml_data":          {},
    "selected_modules":  set(),
    "requirements_list": [],
    "last_run_ts":       None,
    "uml_diagram_png":   None,
}
for _k, _v in _defaults.items():
    if _k not in st.session_state:
        st.session_state[_k] = _v

# ──────────────────────────────────────────────────────────────────────────────
# CONSTANTS
# ──────────────────────────────────────────────────────────────────────────────
AMBIGUITY_MODEL_PATH      = "models/ambiguity"
CLASSIFICATION_MODEL_PATH = "models/classification_final"

VAGUE_TERMS = [
    "fast","quick","slow","rapid","swift","efficient","effective",
    "optimal","optimized","good","bad","better","best","worse","worst",
    "high","low","superior","inferior","robust","reliable","stable",
    "user-friendly","easy","simple","intuitive","convenient","accessible",
    "straightforward","secure","safe","protected","flexible","scalable",
    "adaptable","extensible","many","few","several","multiple","large",
    "small","enough","sufficient","appropriate","suitable","adequate",
    "improved","enhanced","upgraded","properly","correctly","accurately",
    "reasonable","acceptable","satisfactory",
]
WEAK_MODALS = ["may","might","could","should","would","can"]

# ──────────────────────────────────────────────────────────────────────────────
# MODEL LOADERS
# ──────────────────────────────────────────────────────────────────────────────
@st.cache_resource(show_spinner=False)
def load_ambiguity_model():
    try:
        tok   = AutoTokenizer.from_pretrained(AMBIGUITY_MODEL_PATH)
        model = AutoModelForSequenceClassification.from_pretrained(
            AMBIGUITY_MODEL_PATH
        )
        dev = torch.device("cuda" if torch.cuda.is_available() else "cpu")
        model.to(dev).float().eval()
        return tok, model, dev
    except Exception as exc:
        st.error(f"❌ Ambiguity model error: {exc}")
        return None, None, None


@st.cache_resource(show_spinner=False)
def load_classification_model():
    if not os.path.exists(CLASSIFICATION_MODEL_PATH):
        return None, None, None, 0.30
    try:
        tok   = AutoTokenizer.from_pretrained(CLASSIFICATION_MODEL_PATH)
        model = AutoModelForSequenceClassification.from_pretrained(
            CLASSIFICATION_MODEL_PATH
        )
        dev = torch.device("cuda" if torch.cuda.is_available() else "cpu")
        model.to(dev).float().eval()
        threshold = 0.30
        meta_path = os.path.join(CLASSIFICATION_MODEL_PATH, "training_meta.json")
        if os.path.exists(meta_path):
            with open(meta_path) as f:
                threshold = json.load(f).get("best_threshold", 0.30)
        return tok, model, dev, threshold
    except Exception as exc:
        st.warning(f"⚠️ Classification model error: {exc}")
        return None, None, None, 0.30


@st.cache_resource(show_spinner=False)
def load_rule_classifier():
    return RuleBasedClassifier()

# ──────────────────────────────────────────────────────────────────────────────
# FILE UTILITIES
# ──────────────────────────────────────────────────────────────────────────────
def extract_text_from_file(uploaded_file) -> str:
    ext = uploaded_file.name.rsplit(".", 1)[-1].lower()
    try:
        if ext == "txt":
            return uploaded_file.read().decode("utf-8")
        elif ext == "docx":
            doc = Document(io.BytesIO(uploaded_file.read()))
            return "\n".join(
                p.text.strip() for p in doc.paragraphs if p.text.strip()
            )
        elif ext == "pdf":
            pages = []
            with pdfplumber.open(io.BytesIO(uploaded_file.read())) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        pages.append(t)
            return "\n".join(pages)
    except Exception as exc:
        st.error(f"❌ File read error: {exc}")
    return ""


def extract_requirements(text: str) -> list[str]:
    if not text or not text.strip():
        return []

    # First try: split on single newlines (most common paste format)
    lines = [line.strip() for line in text.strip().splitlines() if line.strip()]

    # If we got multiple lines, use them directly
    if len(lines) > 1:
        cleaned = [l for l in lines if len(l) >= 15]
        if cleaned:
            return cleaned

    # Fallback: split on sentence boundaries, double newlines, or
    # requirement prefixes like REQ-001:, FR-001:, 1:, 1., etc.
    reqs = re.split(
        r'(?<=[.!?])\s+(?=[A-Z])|'
        r'\n\n+|'
        r'^(?:REQ|FR|NFR|SRS|RS|UC|US)[-\s]*\d+\s*[:.]\s*|'
        r'(?:^|\n)\s*\d+[:.]\s+',
        text.strip(),
        flags=re.MULTILINE,
    )
    cleaned = [r.strip() for r in reqs if len(r.strip()) >= 15]
    return cleaned if cleaned else [text.strip()]

# ──────────────────────────────────────────────────────────────────────────────
# ANALYSIS HELPERS
# ──────────────────────────────────────────────────────────────────────────────
def detect_vague_terms(text: str):
    if not text:
        return [], []
    low = text.lower()
    vague  = sorted({
        t for t in VAGUE_TERMS
        if re.search(r'\b' + re.escape(t) + r'(?:ly)?\b', low)
    })
    modals = sorted({
        m for m in WEAK_MODALS
        if re.search(r'\b' + re.escape(m) + r'\b', low)
    })
    return vague, modals


def get_ambiguity_reason(text, vague, modals, status):
    if status == "Clear":
        return "✅ Specific and measurable"
    reasons = []
    if vague:  reasons.append(f"Vague: {', '.join(vague[:2])}")
    if modals: reasons.append(f"Weak modals: {', '.join(modals[:2])}")
    has_num  = bool(re.search(r'\d+', text))
    has_unit = bool(re.search(
        r'\b(seconds?|ms|%|users?|MB|GB|minutes?|hours?)\b', text, re.I
    ))
    if not (has_num or has_unit):
        reasons.append("No measurable criteria")
    return " | ".join(reasons) if reasons else "Missing acceptance criteria"


def calculate_rqi(prob_ambiguous: float, text: str) -> float:
    clarity      = round(1 - prob_ambiguous, 3)
    has_num      = bool(re.search(r'\d+', text))
    has_unit     = bool(re.search(r'\b(seconds?|ms|%|users?|MB|GB)\b', text, re.I))
    specificity  = round(min((int(has_num) + int(has_unit)) / 2, 1.0), 3)
    completeness = round(min(len(text.strip().split()) / 15, 1.0), 3)
    return round(0.5*clarity + 0.25*completeness + 0.25*specificity, 3)


def predict_ambiguity(text, tokenizer, model, device):
    if not (tokenizer and model):
        return None
    try:
        inputs = tokenizer(
            text, return_tensors="pt",
            truncation=True, max_length=256, padding=True
        ).to(device)
        with torch.no_grad():
            probs = torch.softmax(model(**inputs).logits.float(), dim=1)
            pred  = torch.argmax(probs, dim=1).item()
        return {
            "label_text":     "Ambiguous" if pred == 1 else "Clear",
            "confidence":     round(probs[0][pred].item(), 4),
            "prob_ambiguous": round(probs[0][1].item(), 4),
        }
    except Exception:
        return None

# ──────────────────────────────────────────────────────────────────────────────
# RULE-BASED FR/NFR CLASSIFIER
# ──────────────────────────────────────────────────────────────────────────────
class RuleBasedClassifier:
    def __init__(self):
        self.fr_verbs = [
            "add","remove","delete","create","edit","update","manage","browse",
            "search","filter","sort","display","show","select","choose","pick",
            "place","submit","upload","download","export","import","login",
            "register","authenticate","view","access","retrieve","store","save",
            "modify","calculate","compute","process","handle","perform","allow",
            "enable","support","provide","facilitate","print","send","receive",
            "transmit","book","reserve","schedule","track","monitor","assign",
            "approve","reject","cancel","register", "login", "log in", "manage",
            "update", "view", "book", "schedule",
            "assign", "record", "generate",
            "track", "create", "submit",
            "download", "access","reset","change","list",
        ]
        self.fr_patterns = [
            r'user.*(?:shall|can|should|must).*(?:add|remove|view|edit|place|select|create)',
            r'(?:system|user).*allows?.*(?:to|users?)',
            r'.*\b(?:add|remove|edit|view|place|select|browse|search|create|manage)\b',
            r'.*\b(?:screen|dialog|form|button|menu|interface|page|dashboard)\b',
        ]
        self.nfr_attrs = [
            "performance","response time","throughput","latency","speed",
            "security","secure","encryption","encrypted","ssl","tls",
            "reliability","reliable","availability","uptime","downtime",
            "scalability","scalable","maintainability","usability",
            "user-friendly","accessibility","compatibility","portability",
            "testability","backup","recovery","resilience","load","concurrent",
            "capacity","efficient","efficiency","optimal","robustness",
            "data integrity","consistency","audit","logging","compliance",
            "privacy","gdpr","memory","cpu","storage","bandwidth",
        ]
        self.nfr_patterns = [
            r'(?:shall|must|should).*(?:respond|return|complete|load).*(?:within|in|less than)\s*\d+',
            r'(?:shall|must).*support.*\d+.*(?:user|concurrent|transaction)',
            r'(?:shall|must).*(?:be|ensure|maintain).*(?:secure|reliable|available|scalable)',
            r'\b(?:throughput|latency|bandwidth|memory|storage|cpu)\b',
            r'.*\bencrypt\b.*',
            r'.*\bbackup\b.*',
            r'.*\b(?:support|handle)\s+\d+\s*(?:user|transaction|request)',
        ]

    def classify(self, text: str) -> dict:
        low = text.lower()
        fr_score = nfr_score = 0.0
        fr_reason = nfr_reason = ""
        for v in self.fr_verbs:
            if re.search(r'\b' + re.escape(v) + r'\b', low):
                fr_score += 1.0
                if not fr_reason:
                    fr_reason = f"Action: '{v}'"
        for p in self.fr_patterns:
            if re.search(p, low):
                fr_score += 2.0
                if not fr_reason:
                    fr_reason = "FR pattern"
        for a in self.nfr_attrs:
            if re.search(r'\b' + re.escape(a) + r'\b', low):
                nfr_score += 1.0
                if not nfr_reason:
                    nfr_reason = f"Attr: '{a}'"
        for p in self.nfr_patterns:
            if re.search(p, low):
                nfr_score += 2.0
                if not nfr_reason:
                    nfr_reason = "NFR pattern"
        total = fr_score + nfr_score
        if total == 0:
            return {
                "label": 0, "label_text": "Functional (FR)",
                "confidence": 0.5,
            }
        if fr_score >= nfr_score:
            return {
                "label": 0, "label_text": "Functional (FR)",
                "confidence": min(fr_score / total, 0.99),
            }
        return {
            "label": 1, "label_text": "Non-Functional (NFR)",
            "confidence": min(nfr_score / total, 0.99),
        }


def predict_fr_nfr_hybrid(text, tok, model, dev, threshold):
    rule_res  = load_rule_classifier().classify(text)
    rule_conf = rule_res["confidence"]
    try:
        inputs = tok(
            text, return_tensors="pt",
            truncation=True, max_length=128, padding=True
        ).to(dev)
        with torch.no_grad():
            probs      = torch.softmax(model(**inputs).logits.float(), dim=1)
            prob_nfr   = probs[0][1].item()
            model_pred = 1 if prob_nfr >= threshold else 0
            model_conf = max(probs[0]).item()
    except Exception:
        model_pred = rule_res["label"]
        model_conf = rule_conf
        prob_nfr   = 0.5

    if rule_conf >= 0.75:
        final_label, final_conf, method = rule_res["label"], rule_conf, "rules"
    else:
        final_label, final_conf, method = model_pred, model_conf, "model"

    return {
        "label_text":  "Non-Functional (NFR)" if final_label == 1 else "Functional (FR)",
        "prob_fr":     round(1 - prob_nfr, 4),
        "prob_nfr":    round(prob_nfr, 4),
        "confidence":  round(final_conf, 4),
        "method_used": method,
    }

# ──────────────────────────────────────────────────────────────────────────────
# CHARTS
# ──────────────────────────────────────────────────────────────────────────────
def make_pie(labels, values, clrs, title):
    fig = go.Figure(go.Pie(
        labels=labels, values=values,
        marker=dict(colors=clrs, line=dict(color="white", width=2)),
        hole=0.4, textinfo="percent+label",
        textfont=dict(size=11),
    ))
    fig.update_layout(
        title=dict(text=title, font=dict(size=13, color="#1e88e5")),
        paper_bgcolor="#f8fafb",
        font=dict(color="#1a237e"),
        height=270,
        margin=dict(t=40, b=10, l=10, r=10),
    )
    return fig

# ──────────────────────────────────────────────────────────────────────────────
# REPORT GENERATORS  (PDF + DOCX embed the PNG diagram)
# ──────────────────────────────────────────────────────────────────────────────
def generate_pdf(results_df: pd.DataFrame, summary: dict,
                 uml_png: bytes = None) -> io.BytesIO:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=landscape(A4),
        rightMargin=25, leftMargin=25, topMargin=25, bottomMargin=25,
    )
    styles = getSampleStyleSheet()
    blue   = colors.HexColor("#1e88e5")
    lblue  = colors.HexColor("#e3f2fd")

    T = ParagraphStyle("T", parent=styles["Heading1"], fontSize=22,
                       textColor=blue, spaceAfter=6,
                       alignment=TA_CENTER, fontName="Helvetica-Bold")
    H = ParagraphStyle("H", parent=styles["Heading2"], fontSize=13,
                       textColor=blue, spaceAfter=8, fontName="Helvetica-Bold")
    C = ParagraphStyle("C", parent=styles["Normal"], fontSize=7.5, leading=10)

    elems = [
        Paragraph("AI-RADA — Requirement Analysis Report", T),
        Paragraph(
            f"Generated: {datetime.datetime.now():%Y-%m-%d %H:%M:%S}  |  "
            f"Requirements analysed: {len(results_df)}",
            styles["Normal"],
        ),
        Spacer(1, 0.2*inch),
        Paragraph("Executive Summary", H),
    ]

    # Summary table
    s_data = [["Metric", "Value"]] + [
        [str(k), str(v)] for k, v in summary.items()
    ]
    s_tbl = Table(s_data, colWidths=[3*inch, 2*inch])
    s_tbl.setStyle(TableStyle([
        ("BACKGROUND",     (0,0),(-1,0), blue),
        ("TEXTCOLOR",      (0,0),(-1,0), colors.white),
        ("FONTNAME",       (0,0),(-1,0), "Helvetica-Bold"),
        ("GRID",           (0,0),(-1,-1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0,1),(-1,-1), [colors.white, lblue]),
        ("FONTSIZE",       (0,0),(-1,-1), 9),
    ]))
    elems += [s_tbl, Spacer(1, 0.2*inch)]

    # UML diagram
    if uml_png:
        elems.append(Paragraph("UML Use Case Diagram", H))
        try:
            elems.append(RLImage(io.BytesIO(uml_png), width=7*inch, height=4.5*inch))
        except Exception:
            pass
        elems.append(Spacer(1, 0.2*inch))

    # Results table
    elems.append(Paragraph("Detailed Results", H))
    cols  = list(results_df.columns)
    c_w   = [0.35*inch if c == "#" else 2.2*inch if c == "Requirement"
             else 1.0*inch for c in cols]
    r_data = [[Paragraph(str(c), C) for c in cols]]
    for _, row in results_df.iterrows():
        r_data.append([
            Paragraph(
                (str(row[c])[:100] + "…") if len(str(row[c])) > 100 else str(row[c]),
                C,
            )
            for c in cols
        ])
    r_tbl = Table(r_data, colWidths=c_w, repeatRows=1)
    r_tbl.setStyle(TableStyle([
        ("BACKGROUND",     (0,0),(-1,0), blue),
        ("TEXTCOLOR",      (0,0),(-1,0), colors.white),
        ("FONTNAME",       (0,0),(-1,0), "Helvetica-Bold"),
        ("GRID",           (0,0),(-1,-1), 0.4, colors.lightgrey),
        ("VALIGN",         (0,0),(-1,-1), "TOP"),
        ("FONTSIZE",       (0,0),(-1,-1), 7.5),
        ("ROWBACKGROUNDS", (0,1),(-1,-1), [colors.white, lblue]),
        ("TOPPADDING",     (0,1),(-1,-1), 3),
        ("BOTTOMPADDING",  (0,1),(-1,-1), 3),
    ]))
    elems.append(r_tbl)
    doc.build(elems)
    buf.seek(0)
    return buf


def generate_docx(results_df: pd.DataFrame, summary: dict,
                  uml_png: bytes = None) -> io.BytesIO:
    doc = DocxDocument()
    doc.core_properties.title  = "AI-RADA Analysis"
    doc.core_properties.author = "AI-RADA"

    h = doc.add_heading("AI-RADA — Requirement Analysis Report", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        f"Generated: {datetime.datetime.now():%Y-%m-%d %H:%M:%S}"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_heading("Executive Summary", 1)
    tbl = doc.add_table(rows=len(summary) + 1, cols=2)
    tbl.style = "Light Grid Accent 1"
    tbl.rows[0].cells[0].text = "Metric"
    tbl.rows[0].cells[1].text = "Value"
    for i, (k, v) in enumerate(summary.items(), 1):
        tbl.rows[i].cells[0].text = str(k)
        tbl.rows[i].cells[1].text = str(v)

    if uml_png:
        doc.add_paragraph()
        doc.add_heading("UML Use Case Diagram", 1)
        try:
            doc.add_picture(io.BytesIO(uml_png), width=Inches(6.0))
        except Exception:
            pass
        doc.add_paragraph()

    doc.add_paragraph()
    doc.add_heading("Detailed Results", 1)
    r_tbl = doc.add_table(rows=len(results_df) + 1, cols=len(results_df.columns))
    r_tbl.style = "Light Grid Accent 1"
    for ci, col in enumerate(results_df.columns):
        r_tbl.rows[0].cells[ci].text = str(col)
    for ri, (_, row) in enumerate(results_df.iterrows(), 1):
        for ci, col in enumerate(results_df.columns):
            r_tbl.rows[ri].cells[ci].text = str(row[col])[:100]

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ──────────────────────────────────────────────────────────────────────────────
# CORE ANALYSIS
# ──────────────────────────────────────────────────────────────────────────────
def analyze_requirements(
    requirements_list, selected_modules,
    amb_tokenizer, amb_model, amb_device,
    clf_tokenizer, clf_model, clf_device, clf_threshold,
):
    results  = []
    uml_data = {}

    total = len(requirements_list)
    prog  = st.progress(0, text="⏳ Starting…")

    for idx, req in enumerate(requirements_list, start=1):
        prog.progress(int(idx / total * 100), text=f"🔍 Analysing {idx}/{total}…")
        item = {"#": idx, "Requirement": req}

        # Validity check first
        validity = predict_validity(req)
        item.update({
            "Validity":     validity["label"],
            "Val. Score":   f"{validity['confidence']:.1%}",
            "Val. Reason":  validity["reason"],
        })

        # Only proceed with other analyses if valid
        if validity["label"] == "Valid Requirement":
            if "ambiguity" in selected_modules and amb_model is not None:
                pred = predict_ambiguity(req, amb_tokenizer, amb_model, amb_device)
                if pred:
                    vague, modals = detect_vague_terms(req)
                    item.update({
                        "Ambiguity":  pred["label_text"],
                        "Amb. Score": f"{pred['confidence']:.1%}",
                        "RQI Score":  calculate_rqi(pred["prob_ambiguous"], req),
                        "Reason":     get_ambiguity_reason(
                            req, vague, modals, pred["label_text"]
                        ),
                    })

            if "classification" in selected_modules and clf_model is not None:
                pred = predict_fr_nfr_hybrid(
                    req, clf_tokenizer, clf_model, clf_device, clf_threshold
                )
                if pred:
                    item.update({
                        "Type":       pred["label_text"],
                        "Type Score": f"{pred['confidence']:.1%}",
                        "P(FR)":      f"{pred['prob_fr']*100:.0f}%",
                        "P(NFR)":     f"{pred['prob_nfr']*100:.0f}%",
                    })
        else:
            # Skip other analyses for invalid requirements
            if "ambiguity" in selected_modules:
                item.update({
                    "Ambiguity":  "N/A",
                    "Amb. Score": "N/A",
                    "RQI Score":  np.nan,
                    "Reason":     "Skipped - Invalid requirement",
                })
            if "classification" in selected_modules:
                item.update({
                    "Type":       "N/A",
                    "Type Score": "N/A",
                    "P(FR)":      "N/A",
                    "P(NFR)":     "N/A",
                })

        results.append(item)

    prog.progress(100, text="✅ Done!")
    time.sleep(0.3)
    prog.empty()

    # ── UML extraction ────────────────────────────────────────────────────────
    if "uml" in selected_modules:
        try:
            parsed   = parse_requirements(requirements_list)
            uml_data = build_graph(parsed)   # now includes actor_groups
            # AI refinement of the extracted graph
            from app.ml.uml.ai_refiner import refine_graph
            with st.spinner("🤖 AI is refining UML extraction..."):
                uml_data = refine_graph(requirements_list, uml_data)
        except Exception as exc:
            st.warning(f"⚠️ UML extraction error: {exc}")

    return results, uml_data

# ══════════════════════════════════════════════════════════════════════════════
# UI — HERO
# ══════════════════════════════════════════════════════════════════════════════
st.markdown("""
<div class="hero">
  <div class="hero-title">🤖 AI-RADA</div>
  <div class="hero-sub">AI-Powered Requirement Analysis &amp; Design Assistant</div>
  <div class="hero-badges">
    <span class="badge">🔍 Ambiguity Detection</span>
    <span class="badge">🏷️ FR / NFR Classification</span>
    <span class="badge">📊 UML Diagrams</span>
    <span class="badge">📈 Quality Metrics</span>
  </div>
</div>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# UI — INPUT PANEL
# ══════════════════════════════════════════════════════════════════════════════
st.markdown('<div class="panel">', unsafe_allow_html=True)
st.markdown('<div class="panel-title">📥 Requirements Input</div>',
            unsafe_allow_html=True)

tab_manual, tab_upload = st.tabs(["✏️ Manual Input", "📁 Upload File"])
requirements_text = ""

with tab_manual:
    requirements_text = st.text_area(
        label="req", label_visibility="collapsed",
        placeholder=(
            "Paste requirements here — one per line or a full SRS document.\n\n"
            "Examples:\n"
            "  REQ-001: The user shall log in with username and password.\n"
            "  REQ-002: The system shall respond within 2 seconds.\n"
            "  REQ-003: Admin manages user accounts and assigns roles.\n"
            "  REQ-004: The customer can place an order and track delivery."
        ),
        height=190,
        key="manual_input",
    )

with tab_upload:
    uploaded_file = st.file_uploader(
        "Choose file:", type=["txt","docx","pdf"],
        label_visibility="collapsed",
    )
    if uploaded_file:
        with st.spinner(f"📄 Reading {uploaded_file.name}…"):
            extracted = extract_text_from_file(uploaded_file)
        if extracted:
            st.success(
                f"✅ Loaded **{uploaded_file.name}** — {len(extracted):,} chars"
            )
            with st.expander("👁️ Preview"):
                st.text(extracted[:600] + ("…" if len(extracted) > 600 else ""))
            requirements_text = extracted
        else:
            st.error("❌ Could not extract text.")

st.markdown("</div>", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# UI — MODULE SELECTOR
# ══════════════════════════════════════════════════════════════════════════════
st.markdown('<div class="panel">', unsafe_allow_html=True)
st.markdown('<div class="panel-title">⚙️ Analysis Modules</div>',
            unsafe_allow_html=True)

col_m1, col_m2, col_m3 = st.columns(3)

with col_m1:
    st.markdown("""
    <div class="module-card">
      <span class="module-icon">🔍</span>
      <div class="module-name">Ambiguity Detection</div>
      <div class="module-desc">DeBERTa model · Vague terms · RQI scoring</div>
    </div>""", unsafe_allow_html=True)
    run_ambiguity = st.checkbox(
        "Ambiguity", value=True, key="cb_amb",
        label_visibility="collapsed",
    )
    st.markdown(
        f"<p style='text-align:center;"
        f"color:{'#43a047' if run_ambiguity else '#bbb'};"
        f"font-size:0.8rem;margin-top:0.3rem;'>"
        f"{'✅ Active' if run_ambiguity else '⬜ Inactive'}</p>",
        unsafe_allow_html=True,
    )

with col_m2:
    st.markdown("""
    <div class="module-card">
      <span class="module-icon">🏷️</span>
      <div class="module-name">FR / NFR Classification</div>
      <div class="module-desc">Hybrid rules + DistilBERT · Confidence scores</div>
    </div>""", unsafe_allow_html=True)
    run_classification = st.checkbox(
        "Classification", value=True, key="cb_clf",
        label_visibility="collapsed",
    )
    st.markdown(
        f"<p style='text-align:center;"
        f"color:{'#43a047' if run_classification else '#bbb'};"
        f"font-size:0.8rem;margin-top:0.3rem;'>"
        f"{'✅ Active' if run_classification else '⬜ Inactive'}</p>",
        unsafe_allow_html=True,
    )

with col_m3:
    st.markdown("""
    <div class="module-card">
      <span class="module-icon">📊</span>
      <div class="module-name">UML Diagrams</div>
      <div class="module-desc">Actor extraction · Clean PNG diagrams</div>
    </div>""", unsafe_allow_html=True)
    run_uml = st.checkbox(
        "UML", value=False, key="cb_uml",
        label_visibility="collapsed",
    )
    st.markdown(
        f"<p style='text-align:center;"
        f"color:{'#43a047' if run_uml else '#bbb'};"
        f"font-size:0.8rem;margin-top:0.3rem;'>"
        f"{'✅ Active' if run_uml else '⬜ Inactive'}</p>",
        unsafe_allow_html=True,
    )

st.divider()

col_btn, col_count, col_tip = st.columns([1.5, 1, 2])

with col_btn:
    execute_btn = st.button(
        "🚀 Analyse Requirements",
        type="primary",
        use_container_width=True,
    )

with col_count:
    n_preview = (
        len(extract_requirements(requirements_text))
        if requirements_text.strip() else 0
    )
    st.markdown(f"""
    <div style='padding:0.8rem;background:#e3f2fd;border-radius:8px;
                text-align:center;border:2px solid #1e88e5;'>
      <span style='font-size:1.6rem;font-weight:900;color:#1e88e5;display:block;'>
        {n_preview}
      </span>
      <span style='font-size:0.75rem;color:#1565c0;'>REQUIREMENTS DETECTED</span>
    </div>""", unsafe_allow_html=True)

with col_tip:
    st.info("💡 Select modules above, then click **Analyse Requirements**.")

st.markdown("</div>", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# EXECUTION
# ══════════════════════════════════════════════════════════════════════════════
if execute_btn:
    if not requirements_text.strip():
        st.error("⚠️ Please enter requirements or upload a file.")
        st.stop()
    if not (run_ambiguity or run_classification or run_uml):
        st.error("⚠️ Please enable at least one module.")
        st.stop()

    selected_modules: set = set()
    if run_ambiguity:      selected_modules.add("ambiguity")
    if run_classification: selected_modules.add("classification")
    if run_uml:            selected_modules.add("uml")

    amb_tokenizer = amb_model = amb_device = None
    clf_tokenizer = clf_model = clf_device = None
    clf_threshold = 0.30

    _ms = st.empty()
    _ms.info("🔄 Loading AI models…")
    if run_ambiguity:
        amb_tokenizer, amb_model, amb_device = load_ambiguity_model()
    if run_classification:
        clf_tokenizer, clf_model, clf_device, clf_threshold = (
            load_classification_model()
        )
    _ms.empty()

    if run_ambiguity and amb_model is None:
        st.error("❌ Ambiguity model failed to load.")
        st.stop()
    if run_classification and clf_model is None:
        st.error("❌ Classification model failed to load.")
        st.stop()

    requirements_list = extract_requirements(requirements_text)

    results, uml_data = analyze_requirements(
        requirements_list, selected_modules,
        amb_tokenizer, amb_model, amb_device,
        clf_tokenizer, clf_model, clf_device, clf_threshold,
    )

    # Generate PNG straight after analysis (ready for PDF/DOCX)
    uml_png = None
    if "uml" in selected_modules and uml_data and uml_data.get("relationships"):
        with st.spinner("🎨 Rendering diagram…"):
            uml_png = draw_uml_diagram_png(uml_data)

    st.session_state.analysis_results  = results
    st.session_state.uml_data          = uml_data
    st.session_state.selected_modules  = selected_modules
    st.session_state.requirements_list = requirements_list
    st.session_state.last_run_ts       = datetime.datetime.now()
    st.session_state.uml_diagram_png   = uml_png

# ══════════════════════════════════════════════════════════════════════════════
# RESULTS
# ══════════════════════════════════════════════════════════════════════════════
if st.session_state.analysis_results is not None:
    results          = st.session_state.analysis_results
    uml_data         = st.session_state.uml_data
    selected_modules = st.session_state.selected_modules
    ts_run           = st.session_state.last_run_ts
    uml_png_data     = st.session_state.uml_diagram_png

    results_df = pd.DataFrame(results)

    st.success(
        f"✅ Analysis complete — **{len(results)}** requirement(s) | "
        f"{ts_run:%Y-%m-%d %H:%M:%S}"
    )

    # ── SUMMARY ──────────────────────────────────────────────────────────────
    st.markdown('<div class="panel">', unsafe_allow_html=True)
    st.markdown('<div class="section-heading">📈 Summary Dashboard</div>',
                unsafe_allow_html=True)

    summary_data: dict = {}
    metric_defs = [(str(len(results)), "Total", "blue")]

    # Validity metrics first
    if "Validity" in results_df.columns:
        n_valid = (results_df["Validity"] == "Valid Requirement").sum()
        n_invalid = len(results_df) - n_valid
        metric_defs += [
            (str(n_valid),   "Valid",   "green"),
            (str(n_invalid), "Invalid", "red"),
        ]
        summary_data.update({
            "Valid Requirements":   int(n_valid),
            "Invalid Requirements": int(n_invalid),
        })

    if "Ambiguity" in results_df.columns:
        n_clear = (results_df["Ambiguity"] == "Clear").sum()
        n_ambig = (results_df["Ambiguity"] == "Ambiguous").sum()
        avg_rqi = pd.to_numeric(results_df["RQI Score"], errors="coerce").mean()
        avg_rqi = float(avg_rqi) if not np.isnan(avg_rqi) else 0.0
        metric_defs += [
            (str(n_clear),     "Clear",     "green"),
            (str(n_ambig),     "Ambiguous", "orange"),
            (f"{avg_rqi:.2f}", "Avg RQI",   "blue"),
        ]
        summary_data.update({
            "Total Requirements":     len(results),
            "Clear Requirements":     int(n_clear),
            "Ambiguous Requirements": int(n_ambig),
            "Average RQI Score":      f"{avg_rqi:.2f}",
        })

    if "Type" in results_df.columns:
        n_fr  = (results_df["Type"] == "Functional (FR)").sum()
        n_nfr = (results_df["Type"] == "Non-Functional (NFR)").sum()
        metric_defs += [
            (str(n_fr),  "Functional",     "blue"),
            (str(n_nfr), "Non-Functional", "orange"),
        ]
        summary_data.update({
            "Functional (FR)":      int(n_fr),
            "Non-Functional (NFR)": int(n_nfr),
        })

    if not summary_data:
        summary_data["Total Requirements"] = len(results)

    st.markdown('<div class="metric-grid">', unsafe_allow_html=True)
    for val, label, color in metric_defs:
        st.markdown(
            f'<div class="metric-card {color}">'
            f'<div class="metric-value">{val}</div>'
            f'<div class="metric-label">{label}</div>'
            f'</div>',
            unsafe_allow_html=True,
        )
    st.markdown("</div></div>", unsafe_allow_html=True)

    # ── CHARTS ────────────────────────────────────────────────────────────────
    charts = []
    if "Validity" in results_df.columns:
        vvc = results_df["Validity"].value_counts()
        charts.append(make_pie(
            list(vvc.index), list(vvc.values),
            ["#43a047","#e53935"], "✅ Validity Distribution",
        ))
    if "Ambiguity" in results_df.columns:
        avc = results_df["Ambiguity"].value_counts()
        charts.append(make_pie(
            list(avc.index), list(avc.values),
            ["#43a047","#f57c00"], "🔍 Ambiguity Distribution",
        ))
    if "Type" in results_df.columns:
        tvc = results_df["Type"].value_counts()
        charts.append(make_pie(
            list(tvc.index), list(tvc.values),
            ["#1e88e5","#f57c00"], "🏷️ FR vs NFR Distribution",
        ))

    if charts:
        st.markdown('<div class="panel">', unsafe_allow_html=True)
        st.markdown('<div class="section-heading">📊 Analytics</div>',
                    unsafe_allow_html=True)
        c_cols = st.columns(len(charts))
        for cc, fig in zip(c_cols, charts):
            with cc:
                st.plotly_chart(fig, use_container_width=True)
        st.markdown("</div>", unsafe_allow_html=True)

    # ── DETAILED RESULTS TABLE ────────────────────────────────────────────────
    st.markdown('<div class="panel">', unsafe_allow_html=True)
    st.markdown('<div class="section-heading">📋 Detailed Results</div>',
                unsafe_allow_html=True)
    st.dataframe(results_df, use_container_width=True,
                 height=400, hide_index=True)
    st.markdown("</div>", unsafe_allow_html=True)

    # ── UML SECTION ───────────────────────────────────────────────────────────
    if "uml" in selected_modules:
        st.markdown('<div class="uml-panel">', unsafe_allow_html=True)
        st.markdown('<div class="uml-title">📊 UML Use Case Diagram</div>',
                    unsafe_allow_html=True)

        if not uml_data or not uml_data.get("relationships"):
            st.warning(
                "⚠️ No relationships could be extracted.  \n"
                "Ensure your requirements mention actors such as "
                "**user, admin, customer, system, librarian, staff**, etc., "
                "and actions such as **login, manage, view, place order**, etc."
            )
        else:
            actors    = uml_data.get("actors",        [])
            use_cases = uml_data.get("use_cases",     [])
            rels      = uml_data.get("relationships", [])
            parsed    = uml_data.get("parsed",        [])

            # Metrics row
            um1, um2, um3 = st.columns(3)
            um1.metric("👤 Actors",    len(actors))
            um2.metric("💼 Use Cases", len(use_cases))
            um3.metric("🔗 Relations", len(rels))

            st.divider()

            # PNG diagram + download button
            if uml_png_data:
                ts_uml = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                st.image(uml_png_data, use_container_width=False)
                st.download_button(
                    label="🖼️ Download PNG",
                    data=uml_png_data,
                    file_name=f"uml_diagram_{ts_uml}.png",
                    mime="image/png",
                    help="150 DPI — ready for reports and presentations",
                )
            else:
                st.info("PNG will appear here after analysis.")

            st.divider()

            # Mapping table
            with st.expander("📋 Actor → Use Case Mapping Table",
                             expanded=False):
                if parsed:
                    map_df = pd.DataFrame([
                        {
                            "#":           p["id"],
                            "Requirement": p["requirement"],
                            "Actor":       p["actor"],
                            "Use Case":    p["action"],
                        }
                        for p in parsed
                    ])
                    st.dataframe(map_df, use_container_width=True,
                                 hide_index=True, height=260)

            # PlantUML source
            with st.expander("📄 PlantUML Source Code", expanded=False):
                from app.ml.uml.plantuml_generator import generate_plantuml
                puml_code = generate_plantuml(uml_data)
                st.code(puml_code, language="text")
                ts_puml = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                st.download_button(
                    label="📥 Download .puml",
                    data=puml_code,
                    file_name=f"usecase_{ts_puml}.puml",
                    mime="text/plain",
                    use_container_width=True,
                )

        st.markdown("</div>", unsafe_allow_html=True)

    # ── EXPORT REPORTS ────────────────────────────────────────────────────────
    st.markdown('<div class="panel">', unsafe_allow_html=True)
    st.markdown('<div class="section-heading">📥 Export Reports</div>',
                unsafe_allow_html=True)

    _png  = uml_png_data if "uml" in selected_modules else None
    ts_dl = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")

    dc1, dc2, dc3 = st.columns(3)

    with dc1:
        st.markdown("""<div class="dl-box">
          <div class="dl-box-title">📄 CSV</div>
          <div class="dl-box-desc">Spreadsheet-ready results table</div>
        </div>""", unsafe_allow_html=True)
        st.download_button(
            "📄 Download CSV",
            data=results_df.to_csv(index=False),
            file_name=f"ai_rada_{ts_dl}.csv",
            mime="text/csv",
            use_container_width=True,
        )

    with dc2:
        st.markdown("""<div class="dl-box">
          <div class="dl-box-title">📕 PDF</div>
          <div class="dl-box-desc">Landscape PDF — includes UML diagram</div>
        </div>""", unsafe_allow_html=True)
        pdf_buf = generate_pdf(results_df, summary_data, _png)
        st.download_button(
            "📕 Download PDF",
            data=pdf_buf,
            file_name=f"ai_rada_{ts_dl}.pdf",
            mime="application/pdf",
            use_container_width=True,
        )

    with dc3:
        st.markdown("""<div class="dl-box">
          <div class="dl-box-title">📘 DOCX</div>
          <div class="dl-box-desc">Word document — includes UML diagram</div>
        </div>""", unsafe_allow_html=True)
        docx_buf = generate_docx(results_df, summary_data, _png)
        st.download_button(
            "📘 Download DOCX",
            data=docx_buf,
            file_name=f"ai_rada_{ts_dl}.docx",
            mime="application/vnd.openxmlformats-officedocument"
                  ".wordprocessingml.document",
            use_container_width=True,
        )

    st.markdown("</div>", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════════════════
st.markdown("""
<div class="footer">
  <strong>🤖 AI-RADA v4.2</strong> — Professional Requirement Analysis &amp; Design Assistant<br>
  <span>Ambiguity Detection · FR/NFR Classification · UML Diagrams ·
        Quality Metrics · Professional Reports</span>
</div>
""", unsafe_allow_html=True)