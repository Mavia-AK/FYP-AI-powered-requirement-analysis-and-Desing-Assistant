import base64
import io
from typing import List, Dict, Any, Optional
from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

# Reportlab for PDF
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

# python-docx for DOCX
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

router = APIRouter()

class ExportRequest(BaseModel):
    results: List[Dict[str, Any]]
    uml_base64: Optional[str] = None

@router.post("/export/pdf")
async def export_pdf(request: ExportRequest):
    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=landscape(A4), rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=30)
        elements = []
        styles = getSampleStyleSheet()

        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#0f172a'),
            spaceAfter=30,
            alignment=1 # Center
        )
        elements.append(Paragraph("AI-RADA: Detailed Analysis Audit", title_style))

        # Table Data
        data = [['UUID', 'Requirement Specification', 'Requirement Type', 'Ambiguity', 'Reason']]
        
        for item in request.results:
            uuid = f"ID-{str(item.get('#', '')).rjust(4, '0')}"
            req_text = Paragraph(item.get('Requirement', ''), styles['Normal'])
            req_type = item.get('Type', 'N/A')
            ambiguity = item.get('Ambiguity', 'N/A')
            reason = Paragraph(item.get('Reason') or item.get('Val. Reason') or 'No reason provided.', styles['Normal'])
            data.append([uuid, req_text, req_type, ambiguity, reason])

        # Table Style
        t = Table(data, colWidths=[60, 250, 100, 80, 250])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0f172a')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('TOPPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f8fafc')),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e2e8f0')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('PADDING', (0, 0), (-1, -1), 8),
        ]))
        elements.append(t)
        elements.append(Spacer(1, 40))

        # UML Diagram
        if request.uml_base64:
            try:
                elements.append(Paragraph("Architectural Model", styles['Heading2']))
                elements.append(Spacer(1, 20))
                
                img_data = base64.b64decode(request.uml_base64)
                img_buffer = io.BytesIO(img_data)
                
                img = Image(img_buffer)
                # Scale image to fit page bounds while maintaining aspect ratio
                max_width = 750
                max_height = 450
                
                ratio_w = max_width / img.drawWidth
                ratio_h = max_height / img.drawHeight
                ratio = min(ratio_w, ratio_h)
                
                if ratio < 1:
                    img.drawWidth = img.drawWidth * ratio
                    img.drawHeight = img.drawHeight * ratio
                
                elements.append(img)
            except Exception as e:
                print(f"Error embedding UML in PDF: {e}")

        doc.build(elements)
        buffer.seek(0)
        
        return StreamingResponse(
            buffer, 
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=AI_RADA_Report.pdf"}
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"PDF Generation failed: {str(e)}")

@router.post("/export/docx")
async def export_docx(request: ExportRequest):
    try:
        doc = Document()
        
        # Title
        heading = doc.add_heading('AI-RADA: Detailed Analysis Audit', 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        # Table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Headers
        hdr_cells = table.rows[0].cells
        headers = ['UUID', 'Requirement Specification', 'Requirement Type', 'Ambiguity', 'Reason']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True

        # Data
        for item in request.results:
            row_cells = table.add_row().cells
            row_cells[0].text = f"ID-{str(item.get('#', '')).rjust(4, '0')}"
            row_cells[1].text = item.get('Requirement', '')
            row_cells[2].text = item.get('Type', 'N/A')
            row_cells[3].text = item.get('Ambiguity', 'N/A')
            row_cells[4].text = item.get('Reason') or item.get('Val. Reason') or 'No reason provided.'

        doc.add_paragraph()

        # UML Diagram
        if request.uml_base64:
            try:
                doc.add_heading('Architectural Model', level=1)
                img_data = base64.b64decode(request.uml_base64)
                
                # Check dimensions to prevent cropping in Word
                from PIL import Image as PILImage
                img_buffer_pil = io.BytesIO(img_data)
                with PILImage.open(img_buffer_pil) as pil_img:
                    w, h = pil_img.size
                    
                img_buffer = io.BytesIO(img_data)
                
                # Max width: 6.5 inches, Max height: 8.5 inches
                aspect_ratio = w / h
                if aspect_ratio < (6.5 / 8.5):
                    # Image is tall, constrain by height
                    doc.add_picture(img_buffer, height=Inches(8.0))
                else:
                    # Image is wide, constrain by width
                    doc.add_picture(img_buffer, width=Inches(6.5))
            except Exception as e:
                print(f"Error embedding UML in DOCX: {e}")

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return StreamingResponse(
            buffer, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=AI_RADA_Report.docx"}
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"DOCX Generation failed: {str(e)}")
